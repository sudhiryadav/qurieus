from fastapi import HTTPException
import uuid
import os
import traceback
from models import Document, DocumentChunk, Embedding, Users
from app.core.config import settings
from sentence_transformers import SentenceTransformer
import fitz
import docx
from app.services.model_health_service import model_health_monitor

try:
    # Initialize embedding model
    embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
    # Start monitoring the model's health
    model_health_monitor.start_monitoring()
except Exception as e:
    print(f"Warning: Could not initialize SentenceTransformer: {str(e)}")
    print("Will use placeholder embeddings if needed")
    embedding_model = None
    # Start monitoring anyway to detect when model becomes available
    model_health_monitor.start_monitoring()

# File storage configuration
UPLOAD_DIR = settings.UPLOAD_DIR
os.makedirs(UPLOAD_DIR, exist_ok=True)

def extract_text_from_pdf(filepath):
    """Extract text from a PDF file."""
    try:
        doc = fitz.open(filepath)
        text = "\n".join([page.get_text() for page in doc])
        return text
    except Exception as e:
        print(f"Error extracting text from PDF {filepath}: {str(e)}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Failed to extract text from PDF: {str(e)}"
        )

def extract_text_from_doc(filepath):
    """Extract text from a DOCX file."""
    try:
        doc = docx.Document(filepath)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Error extracting text from DOCX {filepath}: {str(e)}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Failed to extract text from DOCX: {str(e)}"
        )

def extract_text_from_txt(filepath):
    """Extract text from a plain text file."""
    try:
        with open(filepath, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        with open(filepath, 'r', encoding='latin-1') as file:
            return file.read()
    except Exception as e:
        print(f"Error extracting text from TXT {filepath}: {str(e)}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Failed to extract text from TXT: {str(e)}"
        )

def process_file(file_content: bytes, file_extension: str, original_filename: str, userId: str, db):
    """Process a document file and store it in the database.
    
    Args:
        file_content: The binary content of the file
        file_extension: The file extension (e.g., '.pdf', '.docx')
        original_filename: The original filename
        userId: The ID of the user uploading the file
        db: Database session
    """
    try:
        # First check if embedding model is available
        if not embedding_model:
            raise HTTPException(
                status_code=503,
                detail="Document processing service is currently unavailable. Please try again later."
            )

        print(f"Processing file: {original_filename} for user: {userId}")
        
        # Normalize extension
        ext = file_extension.lower()
        
        # Extract text based on file type
        if ext == '.pdf':
            # Create a temporary BytesIO object for PyMuPDF
            import io
            pdf_file = io.BytesIO(file_content)
            doc = fitz.open(stream=pdf_file, filetype="pdf")
            text = "\n".join([page.get_text() for page in doc])
        elif ext in ['.doc', '.docx']:
            # Create a temporary BytesIO object for python-docx
            import io
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif ext in ['.txt', '.md', '.csv']:
            # Decode text content
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                text = file_content.decode('latin-1')
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file format: {ext}")
        
        # Start a new transaction
        db.begin_nested()
        try:
            # Check if user exists
            user = db.query(Users).filter(Users.id == userId).first()
            if not user:
                raise HTTPException(
                    status_code=404,
                    detail=f"User with ID {userId} not found"
                )
            
            # Generate a unique ID for the document
            file_id = str(uuid.uuid4())
            
            # Split into chunks first
            chunk_size = 1000
            chunk_overlap = 100
            chunks = []
            
            # Create chunks
            for i in range(0, len(text), chunk_size - chunk_overlap):
                chunk_text = text[i:i+chunk_size]
                chunks.append(chunk_text)
            
            print(f"Created {len(chunks)} chunks")

            # Generate embeddings for all chunks
            try:
                embeddings = embedding_model.encode(chunks)
                print(f"Generated embeddings for {len(chunks)} chunks")
            except Exception as e:
                print(f"Error generating embeddings: {str(e)}")
                traceback.print_exc()
                raise HTTPException(
                    status_code=500,
                    detail="Failed to generate document embeddings. Please try again."
                )

            # Create document record
            db_document = Document(
                id=file_id,
                userId=userId,
                fileName=original_filename,
                originalName=original_filename,
                fileType=ext,
                fileSize=len(file_content),  # Use actual file size from content
                content=text
            )
            db.add(db_document)
            
            print(f"Created document record with ID: {file_id}")

            # Create chunks and their embeddings
            for i, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
                # Create chunk
                chunk = DocumentChunk(
                    documentId=file_id,
                    content=chunk_text,
                    chunkIndex=i
                )
                db.add(chunk)
                db.flush()  # Flush to get the chunk ID
                
                # Create embedding
                db_embedding = Embedding(
                    userId=userId,
                    chunkId=chunk.id,
                    vector=embedding.tolist()
                )
                db.add(db_embedding)
            
            # Commit the entire transaction
            db.commit()
            return {
                "message": "Document processed successfully",
                "document_id": file_id,
                "chunks": len(chunks)
            }
            
        except Exception as inner_e:
            # Rollback the nested transaction
            db.rollback()
            raise inner_e
            
    except Exception as e:
        print(f"Error processing file: {str(e)}")
        traceback.print_exc()
        # Make sure to roll back any uncommitted changes
        if db.in_transaction():
            db.rollback()
        raise HTTPException(
            status_code=500, 
            detail=f"Failed to process file: {str(e)}"
        )